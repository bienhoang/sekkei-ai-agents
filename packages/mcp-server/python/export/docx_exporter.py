"""Markdown -> Word (.docx) exporter with cover page, TOC, and JP font support."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import yaml

FRONTMATTER_RE = re.compile(r"^---\n([\s\S]*?)\n---\n([\s\S]*)$")
TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")

JAPANESE_FONT = "MS Mincho"
FALLBACK_FONT = "Noto Serif CJK JP"


def _set_japanese_font(doc: Document) -> None:
    """Apply Japanese font to default styles."""
    for style in doc.styles:
        try:
            if hasattr(style, "font") and style.font is not None:
                style.font.name = JAPANESE_FONT
                rpr = style.element.get_or_add_rPr()
                rpr.set(qn("w:eastAsia"), JAPANESE_FONT)
        except (AttributeError, TypeError):
            continue


def _create_cover(doc: Document, meta: dict, project_name: str) -> None:
    """Add cover page with project info."""
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    title = doc.add_heading(meta.get("doc_type", "設計書"), level=0)
    title.alignment = 1  # center

    doc.add_paragraph()
    info = [
        f"プロジェクト名: {project_name}",
        f"文書種別: {meta.get('doc_type', '')}",
        f"版数: {meta.get('version', '1.0')}",
        f"言語: {meta.get('language', 'ja')}",
    ]
    for line in info:
        p = doc.add_paragraph(line)
        p.alignment = 1  # center

    doc.add_page_break()


def _create_toc(doc: Document) -> None:
    """Insert TOC field — Word renders on open (Ctrl+A → F9)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instr_text)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_char_end)

    doc.add_page_break()


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            if i < len(row.cells):
                row.cells[i].text = val
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def _build_body(doc: Document, content: str) -> None:
    """Parse markdown headings, paragraphs, and tables into Word elements."""
    table_header = None
    table_rows = []

    def flush_table():
        nonlocal table_header, table_rows
        if table_header:
            _add_table(doc, table_header, table_rows)
            table_header = None
            table_rows = []

    for line in content.split("\n"):
        stripped = line.strip()

        # Skip HTML comments and code fences
        if stripped.startswith("<!--") or stripped.startswith("```"):
            flush_table()
            continue

        # Headings
        h_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if h_match:
            flush_table()
            level = min(len(h_match.group(1)), 4)
            doc.add_heading(h_match.group(2).strip(), level=level)
            continue

        # Table separator — skip
        if TABLE_SEP_RE.match(stripped):
            continue

        # Table rows
        row_match = TABLE_ROW_RE.match(stripped)
        if row_match:
            cells = [c.strip() for c in row_match.group(1).split("|")]
            if table_header is None:
                table_header = cells
            else:
                table_rows.append(cells)
            continue

        # Non-table line ends any open table
        flush_table()

        # Non-empty paragraph
        if stripped:
            doc.add_paragraph(stripped)

    flush_table()


def export(content: str, doc_type: str, output_path: str, project_name: str = "") -> dict:
    """Main export entry point: Markdown -> Word (.docx)."""
    meta = {}
    body = content

    fm_match = FRONTMATTER_RE.match(content)
    if fm_match:
        meta = yaml.safe_load(fm_match.group(1)) or {}
        body = fm_match.group(2)

    meta.setdefault("doc_type", doc_type)

    doc = Document()
    _set_japanese_font(doc)
    _create_cover(doc, meta, project_name)
    _create_toc(doc)
    _build_body(doc, body)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    file_size = Path(output_path).stat().st_size
    return {"success": True, "file_path": output_path, "file_size": file_size}
